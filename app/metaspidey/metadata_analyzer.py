import os
import json
import mimetypes
from datetime import datetime
import hashlib
import math
import struct
import zipfile
from pathlib import Path
from collections import Counter

# Import libraries with graceful fallbacks
try:
    from PIL import Image
    from PIL.ExifTags import TAGS as EXIF_TAGS, GPSTAGS
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from mutagen import File as MutagenFile
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False

try:
    import exifread
    EXIFREAD_AVAILABLE = True
except ImportError:
    EXIFREAD_AVAILABLE = False


class MetadataAnalyzer:
    """Advanced Metadata Analyzer for comprehensive file analysis"""
    
    def __init__(self):
        self.file_signatures = {
            b'%PDF': 'application/pdf',
            b'\x89PNG': 'image/png',
            b'\xFF\xD8\xFF': 'image/jpeg',
            b'GIF87a': 'image/gif',
            b'GIF89a': 'image/gif',
            b'PK\x03\x04': 'application/zip',
            b'BM': 'image/bmp',
            b'ID3': 'audio/mpeg',
            b'\x1A\x45\xDF\xA3': 'video/webm',
            b'ftypmp4': 'video/mp4',
            b'OggS': 'audio/ogg',
        }
        
        self.threat_indicators = [
            b'eval(',
            b'base64_decode',
            b'exec(',
            b'shell_exec',
            b'<script',
            b'javascript:',
            b'vbscript:',
            b'cmd.exe',
            b'powershell',
        ]

    def analyze_file(self, filepath, extract_exif=True, calculate_hashes=True, deep_analysis=False):
        """
        Comprehensive file analysis with specialized extractors
        """
        try:
            if not os.path.exists(filepath):
                return {'error': f'File not found: {filepath}'}
            
            print(f"Starting comprehensive analysis of: {filepath}")
            
            # Initialize result structure
            result = {
                'analysis_timestamp': datetime.now().isoformat(),
                'filename': os.path.basename(filepath),
                'file_path': filepath,
                'analysis_version': '2.0'
            }
            
            # Basic metadata (always performed)
            result['basic_info'] = self._extract_basic_info(filepath)
            
            # File signature analysis
            result['signature_analysis'] = self._analyze_file_signature(filepath)
            
            # Hash calculation
            if calculate_hashes:
                result['hashes'] = self._calculate_hashes(filepath)
            
            # Security analysis
            result['security_analysis'] = self._perform_security_analysis(filepath)
            
            # Type-specific analysis
            mime_type = result['basic_info'].get('mime_type', '')
            file_ext = result['basic_info'].get('extension', '').lower()
            
            # Image analysis
            if self._is_image_file(mime_type, file_ext):
                result['image_metadata'] = self._extract_image_metadata(filepath, extract_exif)
            
            # Document analysis
            elif self._is_document_file(mime_type, file_ext):
                result['document_metadata'] = self._extract_document_metadata(filepath, file_ext)
            
            # Audio/Video analysis
            elif self._is_media_file(mime_type, file_ext):
                result['media_metadata'] = self._extract_media_metadata(filepath)
            
            # Archive analysis
            elif self._is_archive_file(mime_type, file_ext):
                result['archive_metadata'] = self._extract_archive_metadata(filepath)
            
            # Deep analysis (entropy, file structure, etc.)
            if deep_analysis:
                result['deep_analysis'] = self._perform_deep_analysis(filepath)
            
            # File categorization
            result['categorization'] = self._categorize_file(result)
            
            print(f"Analysis completed for: {filepath}")
            return result
            
        except Exception as e:
            return {
                'error': f'Analysis failed: {str(e)}',
                'filename': os.path.basename(filepath) if filepath else 'unknown',
                'analysis_timestamp': datetime.now().isoformat(),
                'analysis_version': '2.0'
            }

    def _extract_basic_info(self, filepath):
        """Extract comprehensive basic file information"""
        try:
            stat = os.stat(filepath)
            file_path = Path(filepath)
            
            # MIME type detection
            mime_type, encoding = mimetypes.guess_type(filepath)
            if not mime_type:
                # Fallback to signature detection
                with open(filepath, 'rb') as f:
                    header = f.read(16)
                    for sig, detected_mime in self.file_signatures.items():
                        if header.startswith(sig):
                            mime_type = detected_mime
                            break
                    
                if not mime_type:
                    mime_type = 'application/octet-stream'
            
            return {
                'filename': file_path.name,
                'file_path': str(file_path.resolve()),
                'file_size': stat.st_size,
                'file_size_human': self._format_file_size(stat.st_size),
                'mime_type': mime_type,
                'encoding': encoding,
                'extension': file_path.suffix.lower(),
                'stem': file_path.stem,
                'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'accessed': datetime.fromtimestamp(stat.st_atime).isoformat(),
                'permissions': oct(stat.st_mode)[-3:],
                'is_hidden': file_path.name.startswith('.'),
                'is_executable': os.access(filepath, os.X_OK),
                'owner_uid': stat.st_uid,
                'group_gid': stat.st_gid,
                'inode': stat.st_ino if hasattr(stat, 'st_ino') else None,
            }
            
        except Exception as e:
            return {'error': f'Failed to extract basic info: {str(e)}'}

    def _analyze_file_signature(self, filepath):
        """Analyze file signature and detect potential mismatches"""
        try:
            with open(filepath, 'rb') as f:
                header = f.read(32)  # Read more bytes for better detection
                
            detected_types = []
            for sig, mime_type in self.file_signatures.items():
                if header.startswith(sig):
                    detected_types.append(mime_type)
            
            extension_mime, _ = mimetypes.guess_type(filepath)
            
            return {
                'file_header': header.hex()[:64],  # First 32 bytes as hex
                'detected_types': detected_types,
                'extension_suggests': extension_mime,
                'signature_mismatch': bool(detected_types and extension_mime and 
                                         extension_mime not in detected_types),
                'header_analysis': self._analyze_header_structure(header)
            }
            
        except Exception as e:
            return {'error': f'Signature analysis failed: {str(e)}'}

    def _analyze_header_structure(self, header):
        """Analyze file header structure for additional insights"""
        analysis = {
            'printable_chars': sum(1 for b in header if 32 <= b <= 126),
            'null_bytes': header.count(0),
            'high_entropy_bytes': sum(1 for b in header if b > 127),
            'header_entropy': self._calculate_entropy(header)
        }
        
        # Check for common patterns
        if b'<' in header and b'>' in header:
            analysis['contains_markup'] = True
        
        if any(pattern in header for pattern in [b'http://', b'https://', b'ftp://']):
            analysis['contains_urls'] = True
            
        return analysis

    def _calculate_hashes(self, filepath):
        """Calculate comprehensive file hashes"""
        hashes = {}
        hash_functions = {
            'MD5': hashlib.md5(),
            'SHA1': hashlib.sha1(),
            'SHA256': hashlib.sha256(),
            'SHA512': hashlib.sha512()
        }
        
        try:
            with open(filepath, 'rb') as f:
                while True:
                    chunk = f.read(65536)  # 64KB chunks for better performance
                    if not chunk:
                        break
                    for hash_func in hash_functions.values():
                        hash_func.update(chunk)
            
            for name, hash_func in hash_functions.items():
                hashes[name] = hash_func.hexdigest()
            
            # Add hash metadata
            hashes['hash_generated_at'] = datetime.now().isoformat()
            
        except Exception as e:
            hashes['error'] = f'Hash calculation failed: {str(e)}'
            
        return hashes

    def _perform_security_analysis(self, filepath):
        """Perform security analysis to detect potential threats"""
        try:
            analysis = {
                'threat_indicators_found': [],
                'suspicious_patterns': [],
                'risk_level': 'low',
                'security_notes': []
            }
            
            # Read file content for analysis
            try:
                with open(filepath, 'rb') as f:
                    content = f.read(min(1024 * 1024, os.path.getsize(filepath)))  # Max 1MB
                    
                # Check for threat indicators
                for indicator in self.threat_indicators:
                    if indicator in content:
                        analysis['threat_indicators_found'].append(indicator.decode('utf-8', errors='ignore'))
                
                # Analyze for suspicious patterns
                if len(analysis['threat_indicators_found']) > 2:
                    analysis['risk_level'] = 'high'
                    analysis['security_notes'].append('Multiple threat indicators detected')
                elif len(analysis['threat_indicators_found']) > 0:
                    analysis['risk_level'] = 'medium'
                    analysis['security_notes'].append('Potential threat indicators found')
                
                # Check for obfuscation
                printable_ratio = sum(1 for b in content[:1024] if 32 <= b <= 126) / min(1024, len(content))
                if printable_ratio < 0.1 and os.path.getsize(filepath) > 1024:
                    analysis['suspicious_patterns'].append('Highly obfuscated content')
                    analysis['risk_level'] = 'medium' if analysis['risk_level'] == 'low' else analysis['risk_level']
                
            except Exception as e:
                analysis['security_notes'].append(f'Security analysis limited: {str(e)}')
            
            return analysis
            
        except Exception as e:
            return {'error': f'Security analysis failed: {str(e)}'}

    def _extract_image_metadata(self, filepath, extract_exif=True):
        """Extract comprehensive image metadata"""
        if not PIL_AVAILABLE:
            return {'error': 'PIL not available for image analysis'}
        
        try:
            with Image.open(filepath) as img:
                metadata = {
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'width': img.size[0],
                    'height': img.size[1],
                    'has_transparency': img.mode in ('RGBA', 'LA') or 'transparency' in img.info,
                    'color_mode': self._analyze_color_mode(img.mode),
                    'estimated_colors': len(img.getcolors(maxcolors=256)) if img.mode in ['P', 'L'] else 'Unknown'
                }
                
                # Image quality estimation
                if img.format == 'JPEG':
                    # Rough quality estimation based on file size vs dimensions
                    pixels = img.size[0] * img.size[1]
                    file_size = os.path.getsize(filepath)
                    bytes_per_pixel = file_size / pixels
                    
                    if bytes_per_pixel > 2:
                        quality_estimate = 'High'
                    elif bytes_per_pixel > 0.5:
                        quality_estimate = 'Medium'
                    else:
                        quality_estimate = 'Low'
                    
                    metadata['estimated_quality'] = quality_estimate
                    metadata['bytes_per_pixel'] = round(bytes_per_pixel, 3)
                
                # EXIF extraction
                if extract_exif:
                    exif_data = {}
                    
                    # Try modern method first
                    if hasattr(img, 'getexif'):
                        exif = img.getexif()
                        for tag_id, value in exif.items():
                            tag = EXIF_TAGS.get(tag_id, f'Tag{tag_id}')
                            exif_data[tag] = self._clean_exif_value(value)
                        
                        # GPS data extraction
                        if 'GPSInfo' in exif_data:
                            gps_info = {}
                            for gps_tag_id, gps_value in exif.get_ifd(0x8825).items():
                                gps_tag = GPSTAGS.get(gps_tag_id, f'GPS{gps_tag_id}')
                                gps_info[gps_tag] = self._clean_exif_value(gps_value)
                            exif_data['GPS_Info'] = gps_info
                    
                    # Fallback to ExifRead for more comprehensive data
                    elif EXIFREAD_AVAILABLE:
                        with open(filepath, 'rb') as f:
                            tags = exifread.process_file(f, details=True)
                            for tag_name, tag_value in tags.items():
                                if not tag_name.startswith('JPEGThumbnail'):
                                    exif_data[tag_name] = str(tag_value)
                    
                    metadata['exif'] = exif_data
                    
                    # Extract meaningful info from EXIF
                    metadata['camera_info'] = self._extract_camera_info(exif_data)
                    metadata['location_info'] = self._extract_location_info(exif_data)
                
                return metadata
                
        except Exception as e:
            return {'error': f'Image analysis failed: {str(e)}'}

    def _extract_document_metadata(self, filepath, file_ext):
        """Extract metadata from various document types"""
        try:
            metadata = {'document_type': 'Unknown'}
            
            if file_ext == '.pdf' and PYPDF2_AVAILABLE:
                metadata.update(self._extract_pdf_metadata(filepath))
            
            elif file_ext in ['.docx', '.doc'] and DOCX_AVAILABLE:
                metadata.update(self._extract_docx_metadata(filepath))
            
            elif file_ext in ['.xlsx', '.xls'] and OPENPYXL_AVAILABLE:
                metadata.update(self._extract_xlsx_metadata(filepath))
            
            else:
                # Generic document analysis
                metadata.update(self._extract_generic_document_metadata(filepath))
            
            return metadata
            
        except Exception as e:
            return {'error': f'Document analysis failed: {str(e)}'}

    def _extract_pdf_metadata(self, filepath):
        """Extract comprehensive PDF metadata"""
        try:
            metadata = {'document_type': 'PDF'}
            
            with open(filepath, 'rb') as file:
                pdf = PyPDF2.PdfReader(file)
                
                metadata['num_pages'] = len(pdf.pages)
                metadata['is_encrypted'] = pdf.is_encrypted
                
                if pdf.metadata:
                    doc_info = {}
                    for key, value in pdf.metadata.items():
                        clean_key = key.replace('/', '').replace(' ', '_').lower()
                        doc_info[clean_key] = str(value) if value else None
                    metadata['document_info'] = doc_info
                
                # Analyze first page for content estimation
                if pdf.pages:
                    try:
                        first_page = pdf.pages[0]
                        text_content = first_page.extract_text()
                        metadata['has_text_content'] = bool(text_content.strip())
                        metadata['estimated_text_length'] = len(text_content)
                        metadata['first_page_preview'] = text_content[:200] + '...' if len(text_content) > 200 else text_content
                    except Exception:
                        metadata['text_extraction_failed'] = True
                
                # Security features
                if hasattr(pdf, 'encrypt_dict'):
                    metadata['has_security'] = True
                
            return metadata
            
        except Exception as e:
            return {'document_type': 'PDF', 'error': f'PDF analysis failed: {str(e)}'}

    def _extract_docx_metadata(self, filepath):
        """Extract Word document metadata"""
        try:
            doc = Document(filepath)
            
            metadata = {
                'document_type': 'Word Document',
                'num_paragraphs': len(doc.paragraphs),
                'num_sections': len(doc.sections)
            }
            
            # Core properties
            core_props = doc.core_properties
            if core_props:
                metadata['core_properties'] = {
                    'author': core_props.author,
                    'title': core_props.title,
                    'subject': core_props.subject,
                    'created': core_props.created.isoformat() if core_props.created else None,
                    'modified': core_props.modified.isoformat() if core_props.modified else None,
                    'last_modified_by': core_props.last_modified_by,
                    'revision': core_props.revision,
                    'version': core_props.version,
                    'keywords': core_props.keywords,
                    'comments': core_props.comments
                }
            
            # Content analysis
            full_text = '\n'.join([para.text for para in doc.paragraphs])
            metadata['estimated_text_length'] = len(full_text)
            metadata['word_count'] = len(full_text.split())
            metadata['has_content'] = bool(full_text.strip())
            
            return metadata
            
        except Exception as e:
            return {'document_type': 'Word Document', 'error': f'DOCX analysis failed: {str(e)}'}

    def _extract_xlsx_metadata(self, filepath):
        """Extract Excel document metadata"""
        try:
            workbook = load_workbook(filepath, read_only=True, data_only=True)
            
            metadata = {
                'document_type': 'Excel Spreadsheet',
                'num_worksheets': len(workbook.sheetnames),
                'worksheet_names': workbook.sheetnames
            }
            
            # Properties
            props = workbook.properties
            if props:
                metadata['properties'] = {
                    'title': props.title,
                    'creator': props.creator,
                    'created': props.created.isoformat() if props.created else None,
                    'modified': props.modified.isoformat() if props.modified else None,
                    'last_modified_by': props.lastModifiedBy,
                    'subject': props.subject,
                    'description': props.description,
                    'keywords': props.keywords,
                    'category': props.category
                }
            
            # Analyze worksheets
            worksheets_info = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_info = {
                    'name': sheet_name,
                    'max_row': sheet.max_row,
                    'max_column': sheet.max_column,
                    'has_data': sheet.max_row > 1 or sheet.max_column > 1
                }
                worksheets_info.append(sheet_info)
            
            metadata['worksheets_info'] = worksheets_info
            
            return metadata
            
        except Exception as e:
            return {'document_type': 'Excel Spreadsheet', 'error': f'XLSX analysis failed: {str(e)}'}

    def _extract_media_metadata(self, filepath):
        """Extract metadata from audio/video files"""
        if not MUTAGEN_AVAILABLE:
            return {'error': 'Mutagen not available for media analysis'}
        
        try:
            audio_file = MutagenFile(filepath)
            if audio_file is None:
                return {'error': 'File not recognized as media file'}
            
            metadata = {
                'media_type': 'Audio' if 'audio' in str(type(audio_file)) else 'Video',
                'format': type(audio_file).__name__,
                'length_seconds': getattr(audio_file.info, 'length', 0),
                'bitrate': getattr(audio_file.info, 'bitrate', 0),
                'sample_rate': getattr(audio_file.info, 'sample_rate', 0),
                'channels': getattr(audio_file.info, 'channels', 0),
            }
            
            # Format duration
            if metadata['length_seconds']:
                duration = int(metadata['length_seconds'])
                minutes, seconds = divmod(duration, 60)
                hours, minutes = divmod(minutes, 60)
                metadata['duration_formatted'] = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
            
            # Extract tags
            if audio_file.tags:
                tags = {}
                for key, value in audio_file.tags.items():
                    if isinstance(value, list):
                        value = ', '.join(str(v) for v in value)
                    tags[key] = str(value)
                metadata['tags'] = tags
                
                # Common tag mapping
                common_tags = {
                    'title': ['TIT2', 'TITLE', '\xa9nam'],
                    'artist': ['TPE1', 'ARTIST', '\xa9ART'],
                    'album': ['TALB', 'ALBUM', '\xa9alb'],
                    'date': ['TDRC', 'DATE', '\xa9day'],
                    'genre': ['TCON', 'GENRE', '\xa9gen']
                }
                
                extracted_tags = {}
                for common_name, possible_keys in common_tags.items():
                    for key in possible_keys:
                        if key in tags:
                            extracted_tags[common_name] = tags[key]
                            break
                
                if extracted_tags:
                    metadata['common_tags'] = extracted_tags
            
            return metadata
            
        except Exception as e:
            return {'error': f'Media analysis failed: {str(e)}'}

    def _extract_archive_metadata(self, filepath):
        """Extract metadata from archive files"""
        try:
            metadata = {'archive_type': 'Unknown'}
            
            if filepath.lower().endswith('.zip'):
                metadata.update(self._analyze_zip_archive(filepath))
            else:
                metadata['archive_type'] = 'Generic Archive'
                metadata['note'] = 'Limited analysis available for this archive type'
            
            return metadata
            
        except Exception as e:
            return {'error': f'Archive analysis failed: {str(e)}'}

    def _analyze_zip_archive(self, filepath):
        """Analyze ZIP archive contents"""
        try:
            with zipfile.ZipFile(filepath, 'r') as zip_file:
                metadata = {
                    'archive_type': 'ZIP',
                    'num_files': len(zip_file.filelist),
                    'compressed_size': sum(info.compress_size for info in zip_file.filelist),
                    'uncompressed_size': sum(info.file_size for info in zip_file.filelist),
                }
                
                # Compression ratio
                if metadata['uncompressed_size'] > 0:
                    ratio = metadata['compressed_size'] / metadata['uncompressed_size']
                    metadata['compression_ratio'] = round(ratio, 3)
                    metadata['compression_percentage'] = round((1 - ratio) * 100, 1)
                
                # Analyze file types in archive
                file_types = Counter()
                file_list = []
                
                for info in zip_file.filelist[:100]:  # Limit to first 100 files
                    if not info.is_dir():
                        ext = os.path.splitext(info.filename)[1].lower()
                        file_types[ext or 'no_extension'] += 1
                        
                        file_list.append({
                            'filename': info.filename,
                            'size': info.file_size,
                            'compressed_size': info.compress_size,
                            'modified': datetime(*info.date_time).isoformat(),
                        })
                
                metadata['file_types_distribution'] = dict(file_types.most_common(10))
                metadata['file_list'] = file_list
                
                return metadata
                
        except Exception as e:
            return {'archive_type': 'ZIP', 'error': f'ZIP analysis failed: {str(e)}'}

    def _perform_deep_analysis(self, filepath):
        """Perform comprehensive deep analysis"""
        try:
            analysis = {}
            file_size = os.path.getsize(filepath)
            
            # Entropy analysis
            with open(filepath, 'rb') as f:
                chunk = f.read(min(65536, file_size))  # Analyze first 64KB
                analysis['entropy'] = self._calculate_entropy(chunk)
                analysis['entropy_analysis'] = self._interpret_entropy(analysis['entropy'])
            
            # Byte distribution analysis
            with open(filepath, 'rb') as f:
                sample = f.read(min(8192, file_size))
                byte_counts = Counter(sample)
                
                analysis['byte_distribution'] = {
                    'unique_bytes': len(byte_counts),
                    'most_frequent_byte': byte_counts.most_common(1)[0] if byte_counts else None,
                    'null_byte_percentage': round((byte_counts.get(0, 0) / len(sample)) * 100, 2) if sample else 0,
                    'printable_percentage': round(sum(count for byte, count in byte_counts.items() 
                                                    if 32 <= byte <= 126) / len(sample) * 100, 2) if sample else 0
                }
            
            # File structure analysis
            analysis['structure_analysis'] = self._analyze_file_structure(filepath)
            
            # Pattern analysis
            analysis['pattern_analysis'] = self._analyze_patterns(filepath)
            
            return analysis
            
        except Exception as e:
            return {'error': f'Deep analysis failed: {str(e)}'}

    def _calculate_entropy(self, data):
        """Calculate Shannon entropy of data"""
        if not data:
            return 0.0
            
        byte_counts = Counter(data)
        data_len = len(data)
        entropy = 0.0
        
        for count in byte_counts.values():
            if count > 0:
                probability = count / data_len
                entropy -= probability * math.log2(probability)
        
        return round(entropy, 4)

    def _interpret_entropy(self, entropy):
        """Interpret entropy value"""
        if entropy < 1.0:
            return 'Very low (highly structured/repetitive)'
        elif entropy < 3.0:
            return 'Low (some structure/patterns)'
        elif entropy < 6.0:
            return 'Medium (mixed content)'
        elif entropy < 7.5:
            return 'High (compressed or encrypted)'
        else:
            return 'Very high (likely encrypted/random)'

    def _analyze_file_structure(self, filepath):
        """Analyze internal file structure"""
        try:
            structure = {}
            file_size = os.path.getsize(filepath)
            
            # Sample different parts of the file
            sample_points = min(10, file_size // 1024)  # Sample every KB up to 10 points
            if sample_points < 2:
                sample_points = 1
            
            entropies = []
            with open(filepath, 'rb') as f:
                for i in range(sample_points):
                    position = (file_size * i) // sample_points
                    f.seek(position)
                    chunk = f.read(1024)
                    if chunk:
                        entropies.append(self._calculate_entropy(chunk))
            
            structure['entropy_variation'] = {
                'min': min(entropies) if entropies else 0,
                'max': max(entropies) if entropies else 0,
                'average': round(sum(entropies) / len(entropies), 4) if entropies else 0,
                'variation': round(max(entropies) - min(entropies), 4) if entropies else 0
            }
            
            # Check for common structures
            with open(filepath, 'rb') as f:
                header = f.read(512)
                f.seek(-min(512, file_size), 2)  # Seek to near end
                footer = f.read(512)
            
            structure['has_structured_header'] = self._entropy_category(self._calculate_entropy(header)) == 'low'
            structure['has_structured_footer'] = self._entropy_category(self._calculate_entropy(footer)) == 'low'
            
            return structure
            
        except Exception as e:
            return {'error': f'Structure analysis failed: {str(e)}'}

    def _entropy_category(self, entropy):
        """Categorize entropy level"""
        if entropy < 3.0:
            return 'low'
        elif entropy < 6.0:
            return 'medium'
        else:
            return 'high'

    def _analyze_patterns(self, filepath):
        """Analyze file for interesting patterns"""
        try:
            patterns = {
                'repeated_sequences': [],
                'interesting_strings': [],
                'url_patterns': [],
                'file_paths': []
            }
            
            with open(filepath, 'rb') as f:
                content = f.read(min(32768, os.path.getsize(filepath)))  # First 32KB
            
            # Look for repeated byte sequences
            for seq_len in [4, 8, 16]:
                sequences = {}
                for i in range(len(content) - seq_len):
                    seq = content[i:i+seq_len]
                    if seq in sequences:
                        sequences[seq] += 1
                    else:
                        sequences[seq] = 1
                
                # Find most repeated sequences
                for seq, count in sorted(sequences.items(), key=lambda x: x[1], reverse=True)[:3]:
                    if count > 2:  # Only sequences that repeat more than twice
                        patterns['repeated_sequences'].append({
                            'sequence': seq.hex(),
                            'length': seq_len,
                            'count': count
                        })
            
            # Look for strings (printable sequences)
            try:
                text_content = content.decode('utf-8', errors='ignore')
                
                # URLs
                import re
                urls = re.findall(r'https?://[^\s<>"]+', text_content)
                patterns['url_patterns'] = urls[:10]  # First 10 URLs
                
                # File paths
                paths = re.findall(r'[A-Za-z]:\\[^<>"|?*\s]+|/[^<>"|?*\s]+', text_content)
                patterns['file_paths'] = paths[:10]  # First 10 paths
                
                # Interesting strings (emails, IPs, etc.)
                emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text_content)
                if emails:
                    patterns['interesting_strings'].extend(emails[:5])
                
                ips = re.findall(r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b', text_content)
                if ips:
                    patterns['interesting_strings'].extend(ips[:5])
                    
            except Exception:
                pass  # Skip string analysis if decoding fails
            
            return patterns
            
        except Exception as e:
            return {'error': f'Pattern analysis failed: {str(e)}'}

    # Helper methods for file type detection
    def _is_image_file(self, mime_type, extension):
        return mime_type.startswith('image/') or extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']

    def _is_document_file(self, mime_type, extension):
        return (mime_type.startswith('application/') and 
                any(doc_type in mime_type for doc_type in ['pdf', 'msword', 'officedocument', 'document']) or
                extension in ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.odt', '.ods', '.odp'])

    def _is_media_file(self, mime_type, extension):
        return (mime_type.startswith(('audio/', 'video/')) or 
                extension in ['.mp3', '.mp4', '.avi', '.mov', '.wav', '.flac', '.ogg', '.mkv', '.webm'])

    def _is_archive_file(self, mime_type, extension):
        return (mime_type in ['application/zip', 'application/x-rar-compressed', 'application/x-tar'] or
                extension in ['.zip', '.rar', '.tar', '.gz', '.7z', '.bz2'])

    def _categorize_file(self, analysis_result):
        """Categorize file based on analysis results"""
        basic_info = analysis_result.get('basic_info', {})
        mime_type = basic_info.get('mime_type', '')
        extension = basic_info.get('extension', '')
        
        # Primary category
        if self._is_image_file(mime_type, extension):
            category = 'Image'
        elif self._is_document_file(mime_type, extension):
            category = 'Document'
        elif self._is_media_file(mime_type, extension):
            category = 'Media'
        elif self._is_archive_file(mime_type, extension):
            category = 'Archive'
        elif mime_type.startswith('text/'):
            category = 'Text'
        elif 'executable' in mime_type or extension in ['.exe', '.dll', '.so']:
            category = 'Executable'
        else:
            category = 'Other'
        
        # Risk assessment
        security = analysis_result.get('security_analysis', {})
        risk_level = security.get('risk_level', 'low')
        
        return {
            'primary_category': category,
            'risk_level': risk_level,
            'is_suspicious': risk_level in ['medium', 'high'],
            'file_class': self._get_file_class(category, risk_level)
        }

    def _get_file_class(self, category, risk_level):
        """Get file classification"""
        if risk_level == 'high':
            return f'Suspicious {category}'
        elif risk_level == 'medium':
            return f'Potentially Risky {category}'
        else:
            return f'Normal {category}'

    # Utility methods
    def _format_file_size(self, size_bytes):
        """Format file size in human readable format"""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} PB"

    def _clean_exif_value(self, value):
        """Clean EXIF values for JSON serialization"""
        if isinstance(value, bytes):
            try:
                return value.decode('utf-8', errors='ignore')
            except:
                return str(value)
        elif isinstance(value, (tuple, list)):
            return [self._clean_exif_value(v) for v in value]
        else:
            return str(value)

    def _analyze_color_mode(self, mode):
        """Analyze image color mode"""
        mode_descriptions = {
            '1': 'Bilevel (1-bit pixels, black and white)',
            'L': 'Grayscale (8-bit pixels)',
            'P': 'Palette mode (8-bit pixels, mapped using color palette)',
            'RGB': 'True color (8-bit pixels, RGB)',
            'RGBA': 'True color with transparency (8-bit pixels, RGB + Alpha)',
            'CMYK': 'Color separation (8-bit pixels, CMYK)',
            'YCbCr': 'Color video format (8-bit pixels, Y, Cb, Cr)',
            'LAB': 'L*a*b* color space',
            'HSV': 'Hue, Saturation, Value color space'
        }
        return mode_descriptions.get(mode, f'Unknown mode: {mode}')

    def _extract_camera_info(self, exif_data):
        """Extract camera information from EXIF data"""
        camera_info = {}
        
        camera_fields = {
            'Make': 'camera_make',
            'Model': 'camera_model',
            'Software': 'software',
            'DateTime': 'date_taken',
            'ExposureTime': 'exposure_time',
            'FNumber': 'aperture',
            'ISO': 'iso',
            'FocalLength': 'focal_length',
            'Flash': 'flash'
        }
        
        for exif_key, info_key in camera_fields.items():
            if exif_key in exif_data:
                camera_info[info_key] = exif_data[exif_key]
        
        return camera_info if camera_info else None

    def _extract_location_info(self, exif_data):
        """Extract location information from EXIF GPS data"""
        if 'GPS_Info' not in exif_data:
            return None
        
        gps_info = exif_data['GPS_Info']
        location = {}
        
        if 'GPSLatitude' in gps_info and 'GPSLatitudeRef' in gps_info:
            location['latitude'] = f"{gps_info['GPSLatitude']} {gps_info['GPSLatitudeRef']}"
        
        if 'GPSLongitude' in gps_info and 'GPSLongitudeRef' in gps_info:
            location['longitude'] = f"{gps_info['GPSLongitude']} {gps_info['GPSLongitudeRef']}"
        
        if 'GPSAltitude' in gps_info:
            location['altitude'] = gps_info['GPSAltitude']
        
        return location if location else None

    def _extract_generic_document_metadata(self, filepath):
        """Generic document analysis for unsupported formats"""
        try:
            with open(filepath, 'rb') as f:
                content = f.read(min(4096, os.path.getsize(filepath)))
            
            # Basic text analysis
            try:
                text_content = content.decode('utf-8', errors='ignore')
                word_count = len(text_content.split())
                line_count = len(text_content.split('\n'))
                
                return {
                    'document_type': 'Text/Generic Document',
                    'estimated_word_count': word_count,
                    'estimated_line_count': line_count,
                    'has_text_content': bool(text_content.strip()),
                    'preview': text_content[:200] + '...' if len(text_content) > 200 else text_content
                }
            except:
                return {
                    'document_type': 'Binary Document',
                    'note': 'Unable to extract text content'
                }
                
        except Exception as e:
            return {'error': f'Generic document analysis failed: {str(e)}'}

    # Directory and batch analysis methods
    def analyze_directory(self, directory_path, **kwargs):
        """Analyze all files in a directory with improved handling"""
        try:
            if not os.path.exists(directory_path):
                return {'error': f'Directory not found: {directory_path}'}
            
            if not os.path.isdir(directory_path):
                return {'error': f'Path is not a directory: {directory_path}'}
            
            results = []
            file_count = 0
            max_files = kwargs.get('max_files', 100)  # Allow configurable limit
            
            print(f"Starting directory analysis: {directory_path}")
            
            # Walk through directory
            for root, dirs, files in os.walk(directory_path):
                for filename in files:
                    if file_count >= max_files:
                        break
                    
                    filepath = os.path.join(root, filename)
                    try:
                        print(f"Analyzing file {file_count + 1}/{max_files}: {filename}")
                        result = self.analyze_file(filepath, **kwargs)
                        result['relative_path'] = os.path.relpath(filepath, directory_path)
                        results.append(result)
                        file_count += 1
                    except Exception as e:
                        results.append({
                            'error': str(e),
                            'filename': filename,
                            'file_path': filepath,
                            'relative_path': os.path.relpath(filepath, directory_path)
                        })
                        file_count += 1
                
                if file_count >= max_files:
                    break
            
            # Generate summary statistics
            summary = self._generate_analysis_summary(results)
            
            return {
                'analysis_type': 'directory',
                'directory_path': directory_path,
                'total_files_analyzed': len(results),
                'max_files_limit': max_files,
                'results': results,
                'summary': summary,
                'timestamp': datetime.now().isoformat(),
                'analysis_version': '2.0'
            }
            
        except Exception as e:
            return {
                'error': f'Directory analysis failed: {str(e)}',
                'directory_path': directory_path,
                'timestamp': datetime.now().isoformat()
            }

    def _generate_analysis_summary(self, results):
        """Generate summary statistics from analysis results"""
        try:
            summary = {
                'file_categories': Counter(),
                'risk_levels': Counter(),
                'file_types': Counter(),
                'total_size': 0,
                'error_count': 0
            }
            
            for result in results:
                if 'error' in result:
                    summary['error_count'] += 1
                    continue
                
                # File categories
                if 'categorization' in result:
                    cat = result['categorization']
                    summary['file_categories'][cat.get('primary_category', 'Unknown')] += 1
                    summary['risk_levels'][cat.get('risk_level', 'unknown')] += 1
                
                # File types
                if 'basic_info' in result:
                    extension = result['basic_info'].get('extension', 'no_extension')
                    summary['file_types'][extension] += 1
                    summary['total_size'] += result['basic_info'].get('file_size', 0)
            
            # Convert counters to dicts and format
            summary['file_categories'] = dict(summary['file_categories'])
            summary['risk_levels'] = dict(summary['risk_levels'])
            summary['file_types'] = dict(summary['file_types'].most_common(10))  # Top 10
            summary['total_size_human'] = self._format_file_size(summary['total_size'])
            
            return summary
            
        except Exception as e:
            return {'error': f'Summary generation failed: {str(e)}'}

    def batch_analyze(self, file_paths, **kwargs):
        """Analyze multiple files with progress tracking"""
        results = []
        
        print(f"Starting batch analysis of {len(file_paths)} files")
        
        for i, filepath in enumerate(file_paths):
            try:
                print(f"Analyzing file {i + 1}/{len(file_paths)}: {os.path.basename(filepath)}")
                result = self.analyze_file(filepath, **kwargs)
                result['batch_index'] = i
                results.append(result)
            except Exception as e:
                results.append({
                    'error': str(e),
                    'filename': os.path.basename(filepath),
                    'file_path': filepath,
                    'batch_index': i
                })
        
        summary = self._generate_analysis_summary(results)
        
        return {
            'analysis_type': 'batch',
            'total_files': len(file_paths),
            'results': results,
            'summary': summary,
            'timestamp': datetime.now().isoformat(),
            'analysis_version': '2.0'
        }